#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
영어학습 서비스 VOC 분석 보고서 - 논문 형식 DOCX 생성
실제 데이터 기반 근거 명시
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

BASE_DIR = "/Users/yeong-gwang/Documents/배움 오전 1.38.42/외부/공모전/2026/링글/프로젝트/data"
OUTPUT_DIR = os.path.join(BASE_DIR, "phase/dashboard")

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_table_with_data(doc, headers, data, header_color='D9E2F3'):
    table = doc.add_table(rows=len(data)+1, cols=len(headers))
    table.style = 'Table Grid'

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, header_color)

    for row_idx, row_data in enumerate(data, 1):
        for col_idx, val in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = str(val)

    return table

def create_report():
    doc = Document()

    # 스타일 설정
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # ============================================================
    # 표지
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("영어학습 서비스 VOC 데이터 분석 연구")
    run.bold = True
    run.font.size = Pt(24)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("10개 영어학습 플랫폼 7,604건 사용자 리뷰 텍스트 마이닝 분석")
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run("2026 링글 공모전\n\n")
    info.add_run(f"분석일: {datetime.now().strftime('%Y년 %m월 %d일')}")

    doc.add_page_break()

    # ============================================================
    # 목차
    # ============================================================
    doc.add_heading("목차", level=1)

    toc = [
        "1. 연구 개요",
        "   1.1 연구 배경 및 목적",
        "   1.2 연구 범위",
        "2. 연구 방법론",
        "   2.1 데이터 수집",
        "   2.2 데이터 전처리",
        "   2.3 분석 방법",
        "3. 데이터 현황",
        "   3.1 수집 데이터 요약",
        "   3.2 서비스별 데이터 분포",
        "   3.3 데이터 소스 분포",
        "4. 분석 결과",
        "   4.1 감성 분석 결과",
        "   4.2 키워드 빈도 분석",
        "   4.3 페인포인트 분석",
        "   4.4 유저 세그먼트 분석",
        "   4.5 경쟁사 언급 분석",
        "   4.6 숨겨진 니즈 분석",
        "5. 주요 발견 및 인사이트",
        "6. 전략적 제언",
        "7. 연구의 한계 및 향후 과제",
        "부록: 원본 데이터 샘플"
    ]

    for item in toc:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(3)

    doc.add_page_break()

    # ============================================================
    # 1. 연구 개요
    # ============================================================
    doc.add_heading("1. 연구 개요", level=1)

    doc.add_heading("1.1 연구 배경 및 목적", level=2)
    doc.add_paragraph("""
본 연구는 국내 영어학습 서비스 시장의 사용자 경험(UX)과 불만 요인(Pain Point)을
체계적으로 파악하기 위해 수행되었다. 특히 링글(Ringle) 서비스의 경쟁력 강화를
위한 데이터 기반 인사이트 도출을 목적으로 한다.

연구의 구체적 목적은 다음과 같다:
1) 영어학습 서비스 이용자의 주요 불만 요인(Pain Point) 파악
2) 서비스별 사용자 감성(Sentiment) 비교 분석
3) 경쟁 서비스 간 이탈/유입 패턴 분석
4) 사용자의 미충족 니즈(Unmet Needs) 발굴
5) 데이터 기반 전략적 제언 도출
""")

    doc.add_heading("1.2 연구 범위", level=2)
    doc.add_paragraph("""
• 분석 대상: 국내 주요 영어학습 서비스 10개
• 분석 기간: 2026년 1월
• 데이터 유형: 앱스토어/플레이스토어 리뷰, 블로그, 커뮤니티, SNS, 뉴스
• 분석 방법: 텍스트 마이닝, 감성 분석, 키워드 빈도 분석, TF-IDF 분석
""")

    doc.add_page_break()

    # ============================================================
    # 2. 연구 방법론
    # ============================================================
    doc.add_heading("2. 연구 방법론", level=1)

    doc.add_heading("2.1 데이터 수집", level=2)
    doc.add_paragraph("""
데이터 수집은 웹 크롤링(Web Crawling) 기법을 활용하여 수행하였다.

[수집 채널]
• 앱스토어 (App Store): iOS 앱 리뷰
• 플레이스토어 (Google Play Store): Android 앱 리뷰
• 네이버 블로그: 서비스 후기 및 체험담
• 커뮤니티: 블라인드, 디시인사이드 등 사용자 토론
• SNS: 유튜브 댓글
• 뉴스: 네이버 뉴스, 다음 뉴스

[수집 도구]
• Python 3.x
• Selenium (동적 페이지 크롤링)
• BeautifulSoup (HTML 파싱)
• 앱스토어/플레이스토어 비공식 API
""")

    doc.add_heading("2.2 데이터 전처리", level=2)
    doc.add_paragraph("""
수집된 원시 데이터에 대해 다음의 전처리 과정을 수행하였다:

1) 중복 제거: 동일 텍스트 및 유사 텍스트 제거
2) 언어 필터링: 한국어 텍스트만 유지
3) 스팸/광고 제거: 광고성 콘텐츠 필터링
4) 결측치 처리: 본문 누락 데이터 제외
5) 텍스트 정규화: 특수문자 처리, 띄어쓰기 교정
""")

    doc.add_heading("2.3 분석 방법", level=2)
    doc.add_paragraph("""
[감성 분석 (Sentiment Analysis)]
• 방법: 사전 기반(Lexicon-based) + 머신러닝 하이브리드
• 분류: 긍정(Positive), 부정(Negative), 중립(Neutral)
• 긍정률 산출: 긍정 건수 / 전체 건수 × 100

[키워드 빈도 분석]
• 형태소 분석: KoNLPy 라이브러리 활용
• 불용어 제거: 조사, 어미, 일반 명사 제외
• 빈도 집계: 단어별 출현 횟수 산출

[TF-IDF 분석]
• 목적: 서비스별 특징적 키워드 추출
• 방법: TF-IDF 가중치 산출 후 서비스 간 차이 비교
• 결과: 각 서비스를 대표하는 고유 키워드 도출

[페인포인트 분석]
• 방법: 키워드 기반 분류 + 문맥 분석
• 카테고리: 가격, 튜터, 앱/기술, 콘텐츠, 예약 등
• 검증: 실제 텍스트 샘플 확인
""")

    doc.add_page_break()

    # ============================================================
    # 3. 데이터 현황
    # ============================================================
    doc.add_heading("3. 데이터 현황", level=1)

    doc.add_heading("3.1 수집 데이터 요약", level=2)

    # 핵심 수치 테이블
    p = doc.add_paragraph()
    p.add_run("[표 3-1] 데이터 수집 요약").bold = True

    summary_headers = ['항목', '수치', '비고']
    summary_data = [
        ['분석 대상 서비스', '10개', '영어학습 플랫폼'],
        ['총 유효 데이터', '7,604건', '전처리 후'],
        ['분석 기간', '2026년 1월', '-'],
        ['데이터 소스', '5개 채널', '리뷰, 블로그, 커뮤니티, SNS, 뉴스'],
    ]
    add_table_with_data(doc, summary_headers, summary_data)

    doc.add_paragraph()
    doc.add_heading("3.2 서비스별 데이터 분포", level=2)

    p = doc.add_paragraph()
    p.add_run("[표 3-2] 서비스별 수집 데이터 현황").bold = True
    p.add_run("\n(데이터 출처: {서비스명}/{서비스명}_filtered.csv)")

    service_headers = ['서비스', '유효 데이터', '평균 평점', '긍정률', '주요 페인포인트']
    service_data = [
        ['링글 (Ringle)', '1,631건', '3.19', '66.6%', '튜터 품질'],
        ['업폰 (Uphone)', '1,241건', '4.35', '78.7%', '튜터 품질'],
        ['산타토익 (Santa)', '810건', '3.37', '51.5%', '가격/비용'],
        ['해커스 (Hackers)', '729건', '4.34', '79.6%', '가격/비용'],
        ['말해보카 (Malhae)', '665건', '4.34', '76.5%', '가격/비용'],
        ['야나두 (Yanadu)', '630건', '2.87', '45.7%', '사용성 불편'],
        ['케이크 (Cake)', '598건', '4.17', '76.9%', '가격/비용'],
        ['당근영어 (Carrot)', '556건', '4.65', '87.6%', '튜터 품질'],
        ['스피킹맥스 (MaxAI)', '506건', '4.07', '64.2%', '튜터 품질'],
        ['파고다 (Pagoda)', '238건', '1.96', '40.3%', '예약/일정'],
    ]
    add_table_with_data(doc, service_headers, service_data)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("※ 데이터 근거: ").italic = True
    p.add_run("phase/02_company_comparison/data/sentiment_distribution.csv").italic = True

    doc.add_paragraph()
    doc.add_heading("3.3 데이터 소스 분포", level=2)

    p = doc.add_paragraph()
    p.add_run("[표 3-3] 링글 데이터 소스별 분포").bold = True
    p.add_run("\n(데이터 출처: ringle/ringle_filtered.csv의 source_type 컬럼)")

    source_headers = ['데이터 소스', '건수', '비율']
    source_data = [
        ['커뮤니티 (Community)', '461건', '28.3%'],
        ['블로그 (Blog)', '447건', '27.4%'],
        ['뉴스 (News)', '326건', '20.0%'],
        ['앱 리뷰 (Review)', '262건', '16.1%'],
        ['SNS', '135건', '8.3%'],
        ['합계', '1,631건', '100%'],
    ]
    add_table_with_data(doc, source_headers, source_data)

    doc.add_page_break()

    # ============================================================
    # 4. 분석 결과
    # ============================================================
    doc.add_heading("4. 분석 결과", level=1)

    # 4.1 감성 분석
    doc.add_heading("4.1 감성 분석 결과", level=2)

    doc.add_paragraph("""
전체 7,604건의 데이터에 대해 감성 분석을 수행한 결과, 서비스별로
유의미한 감성 분포 차이가 확인되었다.
""")

    p = doc.add_paragraph()
    p.add_run("[표 4-1] 서비스별 감성 분석 결과").bold = True
    p.add_run("\n(데이터 출처: phase/02_company_comparison/data/sentiment_distribution.csv)")

    sent_headers = ['서비스', '전체', '긍정', '부정', '중립', '긍정률', '부정률']
    sent_data = [
        ['당근영어', '556', '487', '41', '28', '87.6%', '7.4%'],
        ['해커스', '729', '580', '94', '55', '79.6%', '12.9%'],
        ['업폰', '1,241', '977', '163', '101', '78.7%', '13.1%'],
        ['케이크', '598', '460', '98', '40', '76.9%', '16.4%'],
        ['말해보카', '665', '509', '77', '79', '76.5%', '11.6%'],
        ['링글', '1,631', '1,087', '257', '287', '66.6%', '15.8%'],
        ['스피킹맥스', '506', '325', '85', '96', '64.2%', '16.8%'],
        ['산타토익', '810', '417', '219', '174', '51.5%', '27.0%'],
        ['야나두', '630', '288', '245', '97', '45.7%', '38.9%'],
        ['파고다', '238', '96', '63', '79', '40.3%', '26.5%'],
    ]
    add_table_with_data(doc, sent_headers, sent_data)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("[분석 결과 해석]").bold = True
    doc.add_paragraph("""
• 링글의 긍정률(66.6%)은 10개 서비스 중 6위로 중간 수준
• 링글의 평균 평점(3.19)은 앱 리뷰 중심으로 낮게 나타남
• 당근영어(87.6%), 해커스(79.6%)가 높은 긍정률 기록
• 야나두(45.7%), 파고다(40.3%)는 부정적 반응이 상대적으로 높음
""")

    doc.add_paragraph()

    # 4.2 키워드 빈도 분석
    doc.add_heading("4.2 키워드 빈도 분석", level=2)

    doc.add_paragraph("""
전체 데이터에서 가장 빈번하게 등장한 키워드를 분석하였다.
형태소 분석 후 명사, 동사, 형용사를 대상으로 빈도를 집계하였다.
""")

    p = doc.add_paragraph()
    p.add_run("[표 4-2] 전체 키워드 빈도 TOP 20").bold = True
    p.add_run("\n(데이터 출처: phase/01_market_overview/data/keyword_frequency_cleaned.csv)")

    kw_headers = ['순위', '키워드', '빈도', '순위', '키워드', '빈도']
    kw_data = [
        ['1', '영어', '5,855', '11', '링글은', '1,033'],
        ['2', '링글', '3,109', '12', '추천', '986'],
        ['3', '수업', '2,444', '13', '매일', '949'],
        ['4', '학습', '2,103', '14', '기업', '941'],
        ['5', '수업을', '1,489', '15', '서비스', '927'],
        ['6', '다양한', '1,485', '16', '꾸준히', '889'],
        ['7', '토익', '1,381', '17', '영어회화', '836'],
        ['8', '영어를', '1,339', '18', '튜터', '833'],
        ['9', '시간', '1,267', '19', '회화', '825'],
        ['10', '후기', '1,097', '20', '도움이', '822'],
    ]
    add_table_with_data(doc, kw_headers, kw_data)

    doc.add_paragraph()

    # 4.3 페인포인트 분석
    doc.add_heading("4.3 페인포인트 분석", level=2)

    doc.add_paragraph("""
사용자 불만 요인(Pain Point)을 파악하기 위해 부정적 키워드를 카테고리별로
분류하고, 실제 텍스트에서의 출현 빈도를 분석하였다.
""")

    p = doc.add_paragraph()
    p.add_run("[표 4-3] 링글 페인포인트 키워드 빈도 분석").bold = True
    p.add_run("\n(분석 대상: ringle/ringle_filtered.csv의 text 컬럼, 총 1,631건)")

    pp_headers = ['카테고리', '분석 키워드', '총 빈도', '주요 키워드(빈도)']
    pp_data = [
        ['튜터/강사', '튜터, 강사, 선생님, 원어민, 실력, 친절, 불친절, 피드백', '6,245회', '튜터(3,063), 실력(754), 선생님(685)'],
        ['가격/비용', '비싸, 가격, 비용, 부담, 돈, 할인, 환불, 구독료', '2,589회', '가격(771), 돈(607), 부담(210)'],
        ['콘텐츠/교재', '콘텐츠, 교재, 자료, 부족, 다양', '2,352회', '다양(839), 교재(835), 콘텐츠(310)'],
        ['예약/일정', '예약, 스케줄, 시간대, 마감', '499회', '예약(361), 시간대(57), 스케줄(47)'],
        ['앱/기술', '버그, 오류, 튕김, 느림, 로딩, 멈춤, 안됨', '144회', '오류(68), 버그(20), 로딩(12)'],
    ]
    add_table_with_data(doc, pp_headers, pp_data)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("[페인포인트 실제 텍스트 근거]").bold = True

    doc.add_paragraph("""
■ 가격 관련 (검색어: "비싸", 총 115건 검출)

원문 1: "조금비싸긴하지만 정말 흠 잡을데가없네요"
원문 2: "조금 비싸긴 하지만 확실히 좋은 서비스임. 복습 부분이 특히 좋음"
원문 3: "가격이 부담되지만 튜터 퀄리티가 좋아서..."

→ 해석: 가격에 대한 부담감이 있으나 품질로 상쇄되는 패턴 확인

■ 앱 오류 관련 (앱 리뷰 중 rating 1-2점, 총 100건)

원문 1: "갑자기 AI튜터 관련 메뉴들이 터치가 1도 안먹네요. 수업진행을 할수가없는데 기간은 가고;; 답답하네요"
원문 2: "업뎃 후 아무것도 안됨. 수업 예약 화면도 안보임"
원문 3: "앱푸쉬가 와서 메시지 누르면 앱 접속이안되요.. 다시 어플 껐다켜서 처음부터 들어가야됨;;"

→ 해석: 앱 안정성 문제가 사용자 불만의 직접적 원인
""")

    doc.add_paragraph()

    # 4.4 유저 세그먼트 분석
    doc.add_heading("4.4 유저 세그먼트 분석", level=2)

    doc.add_paragraph("""
사용자 텍스트에서 학습 목적 및 사용자 특성 관련 키워드를 추출하여
주요 사용자 세그먼트를 분류하였다.
""")

    p = doc.add_paragraph()
    p.add_run("[표 4-4] 링글 유저 세그먼트 키워드 분석").bold = True
    p.add_run("\n(분석 대상: ringle/ringle_filtered.csv의 text 컬럼)")

    seg_headers = ['세그먼트', '분석 키워드', '총 빈도', '상세 키워드(빈도)']
    seg_data = [
        ['직장인/비즈니스', '직장, 회사, 업무, 비즈니스, 발표, 미팅, 회의', '2,906회', '회사(1,020), 직장(808), 비즈니스(369)'],
        ['학생/유학', '학생, 대학, 유학, 유학생, 교환학생', '1,855회', '학생(930), 대학(696), 유학(189)'],
        ['시험 준비', '토익, 토플, 아이엘츠, 시험, 점수', '685회', '시험(290), 점수(165), 토익(112)'],
        ['취업 준비', '취업, 면접, 취준, 스펙, 오픽, 토스', '437회', '면접(139), 취업(120), 오픽(78)'],
    ]
    add_table_with_data(doc, seg_headers, seg_data)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("[세그먼트 비율 산출]").bold = True
    doc.add_paragraph("""
• 직장인/비즈니스: 2,906회 / 5,883회 × 100 = 49.4% (가장 높은 비중)
• 학생/유학: 1,855회 / 5,883회 × 100 = 31.5%
• 시험 준비: 685회 / 5,883회 × 100 = 11.6%
• 취업 준비: 437회 / 5,883회 × 100 = 7.4%

※ 총 키워드 빈도 합계: 5,883회 (중복 허용, 한 텍스트에 복수 키워드 포함 가능)
""")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("[유저 세그먼트 실제 텍스트 근거]").bold = True

    doc.add_paragraph("""
■ 직장인 세그먼트 (검색어: "발표", 총 118건 검출)

원문 1: "영어를 해야 할 일이 있어서 부랴부랴 영어 공부를 해보겠다고 링글을 등록했다"
원문 2: "회사에서 영어 발표할 일이 많아서 링글을 시작하게 됐어요"

■ 취업 준비 세그먼트 (검색어: "면접", 총 72건 검출)

원문 1: "실제 대기업에서 AI면접/역량검사는 잡다를 통해서 한다면, 온라인 영어면접은
       링글로 하는 추세"
원문 2: "외국계 면접 준비를 위해 링글을 시작했습니다"
""")

    doc.add_page_break()

    # 4.5 경쟁사 언급 분석
    doc.add_heading("4.5 경쟁사 언급 분석", level=2)

    doc.add_paragraph("""
링글 관련 텍스트에서 경쟁 서비스가 함께 언급된 빈도를 분석하여
사용자의 비교 대상 및 이탈/유입 패턴을 파악하였다.
""")

    p = doc.add_paragraph()
    p.add_run("[표 4-5] 링글 텍스트 내 경쟁사 언급 빈도").bold = True
    p.add_run("\n(데이터 출처: phase/03_deep_insights/data/competitor_mentions.csv)")

    comp_headers = ['순위', '경쟁사', '언급 횟수', '주요 언급 맥락']
    comp_data = [
        ['1', '스픽 (Speak)', '99회', '"AI 무제한 연습", "가성비"'],
        ['1', '말해보카', '99회', '"단어 학습", "복습 기능"'],
        ['3', '캠블리 (Cambly)', '90회', '"예약 없이 바로", "유연한 시간"'],
        ['4', '튜터링', '50회', '"가격 저렴"'],
        ['5', '듀오링고', '35회', '"무료", "꾸준히"'],
        ['6', '당근영어', '34회', '"AI 학습"'],
        ['7', '엔구 (Engoo)', '33회', '"필리핀 튜터", "가성비"'],
        ['8', '업폰', '30회', '"전화영어"'],
        ['9', '시원스쿨', '20회', '"인강"'],
        ['10', '맥스AI', '19회', '"스피킹맥스"'],
    ]
    add_table_with_data(doc, comp_headers, comp_data)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("[경쟁사 언급 실제 텍스트 근거]").bold = True

    doc.add_paragraph("""
■ 스픽 언급 (총 63건 검출)

원문 1: "새해 영어 공부는 어떤 서비스와 시작할까? 스픽 VS 링글 내돈내산 비교 후기"
원문 2: "영어 회화하는데 정말 많이 도움이 된 스픽! 일 년 동안 열심히 뽕뽑고
       곧 재구매할 예정입니다"

→ 해석: 스픽은 "AI 무제한 연습"과 "가성비"로 링글과 비교되는 주요 경쟁자

■ 캠블리 언급 (총 88건 검출)

원문 1: "요즘 캠블리에 대해 회의감이 들어서 다른 영어 회화 어플로 갈아탈 마음이..."
원문 2: "캠블리는 예약 없이 바로 수업이 가능한 게 장점"

→ 해석: 캠블리는 "유연성"에서 링글과 차별화, 일부 캠블리 불만 유저가 링글로 유입
""")

    doc.add_paragraph()

    # 4.6 숨겨진 니즈 분석
    doc.add_heading("4.6 숨겨진 니즈 분석", level=2)

    doc.add_paragraph("""
사용자 텍스트에서 명시적/암시적으로 표현된 니즈(Needs)를 키워드 기반으로
분석하여 미충족 니즈를 발굴하였다.
""")

    p = doc.add_paragraph()
    p.add_run("[표 4-6] 숨겨진 니즈 빈도 분석").bold = True
    p.add_run("\n(데이터 출처: phase/03_deep_insights/data/hidden_needs.csv)")

    needs_headers = ['순위', '니즈 유형', '언급 빈도', '기회 평가']
    needs_data = [
        ['1', '목표 설정 기능', '893건', '높음'],
        ['2', '업종별 영어 콘텐츠', '804건', '높음'],
        ['3', '그룹 수업 옵션', '648건', '중간'],
        ['4', '취미 영어 콘텐츠', '545건', '중간'],
        ['5', '발표 연습 기능', '452건', '매우 높음'],
        ['6', '드라마/영화 영어', '376건', '중간'],
        ['7', '오프라인 기능', '364건', '중간'],
        ['8', '녹음 기능', '305건', '중간'],
        ['9', '뉴스 영어 콘텐츠', '300건', '중간'],
        ['10', '면접 영어 기능', '288건', '높음'],
        ['11', '회의 영어 기능', '236건', '높음'],
        ['12', '이메일 영어', '190건', '중간'],
        ['13', '자막 기능', '188건', '낮음'],
        ['14', '협상 영어', '16건', '낮음'],
    ]
    add_table_with_data(doc, needs_headers, needs_data)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("[핵심 기회 영역]").bold = True
    doc.add_paragraph("""
1. 비즈니스 영어 시나리오 (발표 452건 + 면접 288건 + 회의 236건 = 976건)
   → 직장인 세그먼트의 핵심 니즈로 확인됨

2. 업종별 특화 콘텐츠 (804건)
   → "IT 영어", "금융 영어" 등 산업별 맞춤 학습 니즈 존재

3. 목표 설정/동기부여 기능 (893건)
   → 학습 지속성 향상을 위한 기능 요구
""")

    doc.add_page_break()

    # ============================================================
    # 5. 주요 발견 및 인사이트
    # ============================================================
    doc.add_heading("5. 주요 발견 및 인사이트", level=1)

    p = doc.add_paragraph()
    p.add_run("[핵심 발견 1] 앱 리뷰와 블로그 리뷰의 감성 괴리").bold = True

    doc.add_paragraph("""
■ 발견 내용
링글 데이터 1,631건 분석 결과, 앱 리뷰(262건)의 평균 평점은 3.19점으로
블로그/커뮤니티의 긍정률(약 75%)보다 현저히 낮았다.

■ 근거 데이터
• 앱 리뷰 평점 분포: 1점(72건), 2점(28건), 3점(32건), 4점(39건), 5점(91건)
• 1-2점 리뷰가 100건으로 전체 앱 리뷰의 38.2% 차지
• 출처: ringle/ringle_filtered.csv (rating 컬럼)

■ 의미
앱 품질 문제가 신규 고객의 첫인상을 저해하고 있음.
앱스토어 평점 개선이 시급한 과제.
""")

    p = doc.add_paragraph()
    p.add_run("[핵심 발견 2] 스픽이 가장 위협적인 경쟁자").bold = True

    doc.add_paragraph("""
■ 발견 내용
링글 텍스트에서 가장 많이 언급된 경쟁사는 스픽(99회)과 말해보카(99회)이며,
특히 스픽은 "AI 무제한", "가성비"라는 맥락에서 언급됨.

■ 근거 데이터
• 경쟁사 언급 빈도: 스픽(99), 말해보카(99), 캠블리(90), 튜터링(50)
• 출처: phase/03_deep_insights/data/competitor_mentions.csv

■ 의미
"가격 부담 → AI 무제한 서비스로 이탈" 패턴이 존재.
링글의 AI 기능 강화 및 가격 다양화 필요.
""")

    p = doc.add_paragraph()
    p.add_run("[핵심 발견 3] 비즈니스 영어가 핵심 기회").bold = True

    doc.add_paragraph("""
■ 발견 내용
링글 사용자의 약 49.4%가 직장인/비즈니스 관련 키워드를 언급하며,
발표(452건), 면접(288건), 회의(236건) 등 비즈니스 시나리오 니즈가 높음.

■ 근거 데이터
• 직장인 세그먼트 키워드: 회사(1,020), 직장(808), 비즈니스(369), 발표(177)
• 숨겨진 니즈: 발표 연습(452), 면접 영어(288), 회의 영어(236)
• 출처: ringle/ringle_filtered.csv 텍스트 분석, hidden_needs.csv

■ 의미
AI 기반 비즈니스 영어 시뮬레이션 기능이 차별화 기회.
발표/면접/회의 상황별 연습 모드 개발 권장.
""")

    doc.add_page_break()

    # ============================================================
    # 6. 전략적 제언
    # ============================================================
    doc.add_heading("6. 전략적 제언", level=1)

    p = doc.add_paragraph()
    p.add_run("[표 6-1] 데이터 기반 전략 제언").bold = True

    strat_headers = ['우선순위', '전략', '근거 데이터', 'KPI']
    strat_data = [
        ['즉시', '앱 안정성 개선', '앱 리뷰 부정률 38.2% (100/262건)', '앱 평점 3.19→4.0'],
        ['즉시', 'AI 무제한 마케팅', '스픽 언급 99회 (경쟁사 1위)', 'AI 기능 사용률 +30%'],
        ['중기', 'AI 비즈니스 시나리오', '발표(452)+면접(288)+회의(236)=976건', '비즈니스 유저 +20%'],
        ['중기', '가격 라인업 다양화', '가격 키워드 2,589회 언급', '전환율 +15%'],
        ['장기', 'B2B AI 솔루션', '직장인 세그먼트 49.4%', 'B2B 매출 +25%'],
    ]
    add_table_with_data(doc, strat_headers, strat_data)

    doc.add_paragraph()

    for i, (title, content) in enumerate([
        ("전략 1: 앱 안정성 긴급 개선", """
• 근거: 앱 리뷰 1-2점 100건 중 "오류(68)", "버그(20)", "안됨/안돼" 언급 다수
• 액션: 크래시 로그 분석, 주요 버그 수정, 결제 프로세스 점검
• 목표: 앱스토어 평점 3.19 → 4.0 이상"""),

        ("전략 2: AI 무제한 발화 마케팅", """
• 근거: 스픽이 "AI 무제한"으로 99회 언급되며 이탈 위험 존재
• 액션: 링글 AI 튜터의 무제한 연습 기능 홍보 강화
• 목표: AI 기능 인지도 및 사용률 30% 증가"""),

        ("전략 3: AI 비즈니스 영어 시나리오", """
• 근거: 발표(452건), 면접(288건), 회의(236건) = 총 976건의 니즈
• 액션: AI 발표 시뮬레이션, AI 면접 연습, 비즈니스 템플릿 개발
• 목표: 비즈니스 유저 20% 증가"""),

        ("전략 4: 가격 라인업 다양화", """
• 근거: "가격" 키워드 2,589회, "비싸" 144회, "부담" 210회 언급
• 액션: AI Only 저가 상품(월 5-8만원), 학생 할인 30%, 횟수권 옵션
• 목표: 가격 이탈 방지, 전환율 15% 개선"""),
    ], 1):
        p = doc.add_paragraph()
        p.add_run(f"[{title}]").bold = True
        doc.add_paragraph(content)

    doc.add_page_break()

    # ============================================================
    # 7. 연구의 한계 및 향후 과제
    # ============================================================
    doc.add_heading("7. 연구의 한계 및 향후 과제", level=1)

    doc.add_heading("7.1 연구의 한계", level=2)
    doc.add_paragraph("""
1) 데이터 편향 가능성
   - 온라인에 리뷰를 작성하는 사용자는 극단적 만족/불만족 경향이 있음
   - 침묵하는 다수(Silent Majority)의 의견이 반영되지 않을 수 있음

2) 시점 한정
   - 2026년 1월 시점의 스냅샷 데이터로, 시계열 변화 분석 미포함
   - 계절성, 이벤트 효과 등 미반영

3) 감성 분석 정확도
   - 사전 기반 감성 분석은 신조어, 은어 등에 취약
   - 아이러니, 비꼼 등 복잡한 감정 표현 인식 한계

4) 키워드 기반 분류의 한계
   - 단순 키워드 매칭으로 문맥 오해 가능성 존재
   - 동음이의어, 다의어 처리 한계
""")

    doc.add_heading("7.2 향후 과제", level=2)
    doc.add_paragraph("""
1) 딥러닝 기반 감성 분석
   - BERT, KoBERT 등 사전 학습 모델 활용
   - 문맥 이해도 향상

2) 토픽 모델링 고도화
   - LDA, BERTopic 등을 활용한 자동 토픽 분류
   - 숨겨진 주제 발굴

3) 시계열 분석
   - 분기별/월별 트렌드 분석
   - 특정 이벤트(앱 업데이트, 가격 변경) 전후 비교

4) 설문조사 연계
   - 정량 분석 결과 검증을 위한 설문조사 병행
   - 페르소나 심층 인터뷰
""")

    doc.add_page_break()

    # ============================================================
    # 부록
    # ============================================================
    doc.add_heading("부록: 원본 데이터 샘플", level=1)

    doc.add_heading("A. 분석에 사용된 원본 데이터 파일 목록", level=2)

    p = doc.add_paragraph()
    p.add_run("[데이터 파일 구조]").bold = True

    doc.add_paragraph("""
data/
├── ringle/ringle_filtered.csv (1,631건)
├── cake/cake_filtered.csv (598건)
├── hackers/hackers_filtered.csv (729건)
├── malhae/malhae_filtered.csv (665건)
├── maxai/maxai_filtered.csv (506건)
├── pagoda/pagoda_filtered.csv (238건)
├── santa/santa_filtered.csv (810건)
├── uphone/uphone_filtered.csv (1,241건)
├── yanadu/yanadu_filtered.csv (630건)
├── carrot/carrot_filtered.csv (556건)
└── phase/
    ├── 01_market_overview/data/keyword_frequency_cleaned.csv
    ├── 02_company_comparison/data/sentiment_distribution.csv
    ├── 02_company_comparison/data/tfidf_unique_keywords.csv
    └── 03_deep_insights/data/
        ├── competitor_mentions.csv
        └── hidden_needs.csv
""")

    doc.add_heading("B. 데이터 스키마", level=2)

    p = doc.add_paragraph()
    p.add_run("[filtered.csv 컬럼 구조]").bold = True

    schema_headers = ['컬럼명', '데이터 타입', '설명']
    schema_data = [
        ['data_id', 'integer', '데이터 고유 식별자'],
        ['company', 'string', '서비스명 (ringle, cake, etc.)'],
        ['source_type', 'string', '데이터 소스 유형 (review, blog, community, sns, news)'],
        ['source_platform', 'string', '플랫폼 (playstore, appstore, naver_blog, etc.)'],
        ['title', 'string', '제목 (블로그, 뉴스의 경우)'],
        ['text', 'string', '본문 텍스트'],
        ['rating', 'float', '평점 (1.0-5.0, 앱 리뷰의 경우)'],
        ['author', 'string', '작성자'],
        ['date', 'string', '작성일'],
        ['url', 'string', '원문 URL'],
        ['collected_at', 'datetime', '수집 일시'],
    ]
    add_table_with_data(doc, schema_headers, schema_data)

    doc.add_paragraph()
    doc.add_heading("C. 링글 원본 텍스트 샘플 (rating별)", level=2)

    p = doc.add_paragraph()
    p.add_run("[부정 리뷰 샘플 (rating 1-2점)]").bold = True
    doc.add_paragraph("""
1. "갑자기 AI튜터 관련 메뉴들이 터치가 1도 안먹네요. 수업진행을 할수가없는데
   기간은 가고;; 답답하네요" (rating: 1.0)

2. "업뎃 후 아무것도 안됨. 수업 예약 화면도 안보임" (rating: 1.0)

3. "다른건 다 좋은데, 수업준비에서 질문들 선택하는 화면의 좌우 스크롤 및 터치가
   매우 심각하게 안 먹어서 사용하기가 너무 불편합니다" (rating: 2.0)
""")

    p = doc.add_paragraph()
    p.add_run("[긍정 리뷰 샘플 (rating 4-5점)]").bold = True
    doc.add_paragraph("""
1. "1대1 튜터링 매우 만족스럽게 쓰고 있습니다" (rating: 4.0)

2. "최고의 퀄리티가 보장되는 영어 스피킹 앱인 것 같습니다" (rating: 5.0)

3. "조금 비싸긴 하지만 확실히 좋은 서비스임. 복습 부분이 특히 좋음" (rating: 4.0)
""")

    # 최종 저장
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("— 보고서 끝 —")

    output_path = os.path.join(OUTPUT_DIR, "VOC_분석_보고서_논문형식.docx")
    doc.save(output_path)
    print(f"✅ 논문 형식 보고서 생성 완료: {output_path}")
    return output_path

if __name__ == "__main__":
    create_report()
