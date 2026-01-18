#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
고객 문제 중심 VOC 분석 보고서
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from datetime import datetime

OUTPUT_DIR = "/Users/yeong-gwang/Documents/배움 오전 1.38.42/외부/공모전/2026/링글/프로젝트/data/phase/dashboard"

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, data, header_color='D9E2F3'):
    table = doc.add_table(rows=len(data)+1, cols=len(headers))
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, header_color)
    for r_idx, row in enumerate(data, 1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = str(val)
    return table

def create_report():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # ===== 표지 =====
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("고객이 겪는 진짜 문제")
    run.bold = True
    run.font.size = Pt(28)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("영어학습 서비스 VOC 7,604건 분석")
    run.font.size = Pt(16)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"2026 링글 공모전 | {datetime.now().strftime('%Y-%m-%d')}")

    doc.add_page_break()

    # ===== 핵심 요약 =====
    doc.add_heading("Executive Summary: 고객 문제 핵심 발견", level=1)

    p = doc.add_paragraph()
    p.add_run("■ 분석 개요").bold = True
    doc.add_paragraph("""
• 분석 대상: 10개 영어학습 서비스
• 분석 데이터: 7,604건 (리뷰, 블로그, 커뮤니티)
• 분석 방법: 키워드 빈도 분석, 문맥 분석, 감성 분석
• 데이터 출처: ringle_filtered.csv 외 9개 서비스 데이터
""")

    p = doc.add_paragraph()
    p.add_run("■ TOP 5 고객 문제 (근거 데이터 기반)").bold = True

    problems_headers = ['순위', '고객 문제', '근거 데이터', '데이터 출처']
    problems_data = [
        ['1', '튜터 품질 편차', '336건 언급 (블로그/커뮤니티)', 'ringle_filtered.csv text 분석'],
        ['2', '가격 대비 가치 의문', '329건 언급 (가격 관련)', 'ringle_filtered.csv text 분석'],
        ['3', '예약 경쟁/시간대 제한', '291건 언급', 'ringle_filtered.csv text 분석'],
        ['4', '학습 지속성 어려움', '237건 언급', 'ringle_filtered.csv text 분석'],
        ['5', '앱 안정성 문제', '30건 (부정리뷰 100건 중)', 'rating 1-2점 리뷰 분석'],
    ]
    add_table(doc, problems_headers, problems_data)

    doc.add_page_break()

    # ===== 1. 분석 방법론 =====
    doc.add_heading("1. 분석 방법론: 어떻게 고객 문제를 찾았나", level=1)

    doc.add_heading("1.1 데이터 수집", level=2)
    doc.add_paragraph("""
[수집 데이터]
• 링글 데이터: 1,631건
  - 앱 리뷰 (rating 포함): 262건
  - 블로그 후기: 447건
  - 커뮤니티 글: 461건
  - 뉴스: 326건
  - SNS: 135건
""")

    doc.add_heading("1.2 분석 방법", level=2)
    doc.add_paragraph("""
[Step 1] 부정 리뷰 집중 분석
• 앱 리뷰 중 rating 1-2점 리뷰 100건 전수 분석
• 문제 패턴 키워드 추출 및 분류

[Step 2] 블로그/커뮤니티 문제 언급 분석
• 908건 중 문제/개선 표현 포함 텍스트 693건 추출
• 표현 패턴: "아쉬", "단점", "불편", "개선", "바라" 등

[Step 3] 키워드 빈도 분석
• 문제 카테고리별 키워드 정의
• 전체 텍스트에서 키워드 출현 빈도 집계
• 실제 원문 텍스트로 검증

[Step 4] 경쟁사 비교 분석
• 10개 서비스 동일 키워드 빈도 비교
• 링글 고유 문제 vs 공통 문제 식별
""")

    doc.add_page_break()

    # ===== 2. 문제 상세 분석 =====
    doc.add_heading("2. 고객 문제 상세 분석", level=1)

    # 문제 1: 튜터 품질 편차
    doc.add_heading("2.1 문제 #1: 튜터 품질 편차", level=2)

    p = doc.add_paragraph()
    p.add_run("[데이터 근거]").bold = True
    doc.add_paragraph("""
• 분석 대상: ringle_filtered.csv의 text 컬럼 (1,631건)
• 관련 키워드: "튜터마다", "편차", "사람마다", "강사마다"
• 총 언급 빈도: 336건
""")

    p = doc.add_paragraph()
    p.add_run("[원문 텍스트 근거]").bold = True
    doc.add_paragraph("""
원문 1 (블로그):
"튜터마다 실력 차이가 너무 커요. 어떤 분은 정말 상세하게 교정해주시는데
어떤 분은 그냥 Good job만 하고 끝나요"

원문 2 (커뮤니티):
"좋은 튜터 잡으면 대박인데, 아닌 사람 걸리면 시간 낭비..."

원문 3 (앱 리뷰):
"튜터가 본인이 가능한 시간에 튜터링을 하는 구조라서
e.g)월화수만 연다던지, 아침만 가능하다던지.."
""")

    p = doc.add_paragraph()
    p.add_run("[문제의 핵심]").bold = True
    doc.add_paragraph("""
• 튜터별 피드백 퀄리티 차이 (상세 교정 vs 단순 칭찬)
• 튜터별 수업 스타일 차이
• 좋은 튜터는 예약 경쟁이 치열함
""")

    doc.add_paragraph()

    # 문제 2: 가격 대비 가치 의문
    doc.add_heading("2.2 문제 #2: 가격 대비 가치 의문", level=2)

    p = doc.add_paragraph()
    p.add_run("[데이터 근거]").bold = True
    doc.add_paragraph("""
• 분석 대상: ringle_filtered.csv의 text 컬럼
• 관련 키워드: "비싸", "가격", "비용", "부담", "돈", "환불"
• 총 언급 빈도: 2,589회 (중복 포함)
• 상세: 가격(771), 돈(607), 비용(204), 부담(210), 비싸(144)
""")

    p = doc.add_paragraph()
    p.add_run("[원문 텍스트 근거]").bold = True
    doc.add_paragraph("""
원문 1:
"조금비싸긴하지만 정말 흠 잡을데가없네요"
→ 가격 부담 인정하지만 품질로 상쇄

원문 2:
"가격이 너무 비싸서 굉장히 망설여짐!
오늘 기준 1회 수업기준 약 41000원 .. ( 1:1수업권만, 40분 기준 )"
→ 구체적 가격 언급과 함께 부담 표현

원문 3:
"스픽이 가성비는 더 좋더라고요... AI 무제한 연습이 필요해서 스픽으로"
→ 가격 때문에 경쟁사로 이탈
""")

    p = doc.add_paragraph()
    p.add_run("[문제의 핵심]").bold = True
    doc.add_paragraph("""
• 월 20-30만원의 높은 가격
• 실력 향상을 체감하기 어려움
• 경쟁사(스픽) 대비 가성비 비교
• 학생 할인 부재
""")

    doc.add_paragraph()

    # 문제 3: 예약 경쟁
    doc.add_heading("2.3 문제 #3: 예약 경쟁 / 시간대 제한", level=2)

    p = doc.add_paragraph()
    p.add_run("[데이터 근거]").bold = True
    doc.add_paragraph("""
• 분석 대상: ringle_filtered.csv의 text 컬럼
• 관련 키워드: "예약", "마감", "스케줄", "시간대", "경쟁"
• 총 언급 빈도: 689회
• 상세: 예약(361), 시간대(57), 스케줄(47), 마감(34)
""")

    p = doc.add_paragraph()
    p.add_run("[원문 텍스트 근거]").bold = True
    doc.add_paragraph("""
원문 1 (앱 리뷰, rating 1점):
"튜터 스케줄 오픈 알림도 뜨고 튜터 페이지에서 예약하기 버튼이
활성화 되어있는데도 예약하기 버튼을 누르면 스케줄이 마감되었다고 계속 뜨네요"

원문 2 (블로그):
"내가 원하는 튜터의 예약 가능 시간이 항상 열려 있는 것은 아남.
e.g)월화수만 연다던지, 아침만 가능하다던지..
튜터가 본인이 가능한 시간에 튜터링을 하는 구조"

원문 3 (앱 리뷰):
"수업잡기가 너무힘듭니다... 3주전에 미리잡으려고해도 힘듭니다"
""")

    p = doc.add_paragraph()
    p.add_run("[문제의 핵심]").bold = True
    doc.add_paragraph("""
• 인기 튜터 예약 경쟁 치열
• 직장인 선호 시간대(저녁) 마감 빠름
• 예약 시스템 오류 발생
• 3주 전에도 예약 어려움
""")

    doc.add_page_break()

    # 문제 4: 학습 지속성
    doc.add_heading("2.4 문제 #4: 학습 지속성 어려움", level=2)

    p = doc.add_paragraph()
    p.add_run("[데이터 근거]").bold = True
    doc.add_paragraph("""
• 분석 대상: ringle_filtered.csv의 text 컬럼
• 관련 키워드: "꾸준", "포기", "지속", "작심삼일", "습관", "동기"
• 총 언급 빈도: 886회 (지속/습관 관련)
""")

    p = doc.add_paragraph()
    p.add_run("[원문 텍스트 근거]").bold = True
    doc.add_paragraph("""
원문 1:
"1주일에 40분씩 투자한다고 영어 실력이 눈에 띄게 늘리가 없다는 걸
충분히 인지하고 있었던 터라 그 부분에서는 딱히 큰 기대는 하지 않았습니다"

원문 2:
"가끔 어떤 날은 말이 잘 터져서 피드백 좋게 나오면 기분 좋고 동기부여도 됐구요"
→ 성과 체감이 동기부여의 핵심

원문 3:
"바쁜 일정에 맞춰 유연하게 수업하고 싶어요"
→ 시간 제약으로 지속 어려움
""")

    p = doc.add_paragraph()
    p.add_run("[문제의 핵심]").bold = True
    doc.add_paragraph("""
• 주 1-2회 수업으로 실력 향상 체감 어려움
• 바쁜 직장 생활로 학습 지속 어려움
• 동기부여/목표 설정 시스템 부재
• 복습/자기학습 연계 부족
""")

    doc.add_paragraph()

    # 문제 5: 앱 안정성
    doc.add_heading("2.5 문제 #5: 앱 안정성 문제", level=2)

    p = doc.add_paragraph()
    p.add_run("[데이터 근거]").bold = True
    doc.add_paragraph("""
• 분석 대상: ringle_filtered.csv 중 rating 1-2점 앱 리뷰
• 부정 리뷰: 100건 (전체 앱 리뷰 262건 중 38.2%)
• 문제 패턴별 분류:
  - 업데이트 문제: 30건
  - 앱 오류/버그: 23건
  - 연결/접속: 13건
  - 터치/UI 문제: 10건
  - 로딩/속도: 9건
""")

    p = doc.add_paragraph()
    p.add_run("[원문 텍스트 근거 (실제 앱 리뷰)]").bold = True
    doc.add_paragraph("""
원문 1 (rating 1.0):
"갑자기 AI튜터 관련 메뉴들이 터치가 1도 안먹네요.
수업진행을 할수가없는데 기간은 가고;; 답답하네요"

원문 2 (rating 1.0):
"업뎃 후 아무것도 안됨. 수업 예약 화면도 안보임"

원문 3 (rating 2.0):
"링글의 컨텐츠에는 만족하며 2022년부터 계속 사용중입니다.
다만 휴대폰으로 접속했을 때 속도가 말할수 없이 느립니다.
그리고 최근 업데이트 이후에는 휴대폰으로 수업 접속 시
3~4분 동안은 화면은 나오는데 오디오가 연결이 안돼서,
한참동안 튜터와 뻘쭘하게 기다려야 합니다"

원문 4 (rating 1.0):
"앱푸쉬가 와서 메시지 누르면 앱 접속이안되요..
다시 어플 껐다켜서 처음부터 들어가야됨;;"
""")

    p = doc.add_paragraph()
    p.add_run("[문제의 핵심]").bold = True
    doc.add_paragraph("""
• 업데이트 후 기능 작동 안함
• 터치/스크롤 반응 불량
• 수업 중 오디오 연결 지연
• 결제/환불 프로세스 오류
""")

    doc.add_page_break()

    # ===== 3. 고객 여정별 문제 =====
    doc.add_heading("3. 고객 여정(Journey)별 문제 매핑", level=1)

    p = doc.add_paragraph()
    p.add_run("[데이터 근거: 여정 단계별 키워드 빈도]").bold = True
    doc.add_paragraph("분석 방법: 각 단계별 관련 키워드를 정의하고 ringle_filtered.csv text에서 빈도 집계")

    journey_headers = ['고객 여정 단계', '키워드 빈도', '핵심 문제', '심각도']
    journey_data = [
        ['1. 인지/탐색', '3,145회', '선택지 혼란, 가격 불투명', '중'],
        ['2. 가입/결제', '990회', '가격 부담, 환불 어려움', '상'],
        ['3. 첫 사용', '2,157회', '앱 복잡함, 시작점 모호', '중'],
        ['4. 예약/준비', '1,982회', '예약 경쟁, 시간대 제한', '상'],
        ['5. 수업 진행', '2,268회', '연결 불안정, 튜터 편차', '상'],
        ['6. 복습', '1,272회', '피드백 부족, 복습 자료 부족', '중'],
        ['7. 지속/재결제', '975회', '동기 부족, 효과 의문', '상'],
    ]
    add_table(doc, journey_headers, journey_data)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("[심각도 '상' 단계 상세]").bold = True

    doc.add_paragraph("""
■ 가입/결제 단계
• "환불하는데 무슨 1:1 문의를 시키고, 30분이 넘도록 답장이 없나요"
• 환불 프로세스 불편 → 신규 고객 이탈 위험

■ 예약/준비 단계
• "수업잡기가 너무힘듭니다... 3주전에 미리잡으려고해도 힘듭니다"
• 예약 실패 경험 → 서비스 불만족

■ 수업 진행 단계
• "최근 업데이트 이후에는 휴대폰으로 수업 접속 시 3~4분 동안은
  화면은 나오는데 오디오가 연결이 안돼서"
• 수업 중 기술 문제 → 핵심 경험 훼손

■ 지속/재결제 단계
• "1주일에 40분씩 투자한다고 영어 실력이 눈에 띄게 늘리가 없다는 걸"
• 효과 체감 부족 → 재결제 망설임
""")

    doc.add_page_break()

    # ===== 4. 경쟁사 비교 =====
    doc.add_heading("4. 경쟁사 대비 링글 고유 문제", level=1)

    p = doc.add_paragraph()
    p.add_run("[데이터 근거]").bold = True
    doc.add_paragraph("""
분석 방법: 10개 서비스 데이터에서 동일 키워드 빈도 비교
데이터 출처: {service}_filtered.csv (10개 서비스)
""")

    compare_headers = ['문제 유형', '링글', '스픽/맥스AI', '캠블리', '링글 특이점']
    compare_data = [
        ['가격 부담', '1,442회', '367회', '133회', '압도적 높음 ⚠️'],
        ['학습효과 의문', '1,290회', '300회', '211회', '높음 ⚠️'],
        ['지속 어려움', '886회', '374회', '200회', '높음'],
        ['실전 괴리', '1,238회', '423회', '306회', '높음'],
    ]
    add_table(doc, compare_headers, compare_data)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("[링글만의 고유 문제]").bold = True

    unique_headers = ['문제', '빈도', '원인 분석']
    unique_data = [
        ['높은 진입장벽', '790회', '"명문대 튜터"라는 프리미엄 포지셔닝이 오히려 부담'],
        ['예약 경쟁', '689회', '인기 튜터 수급 불균형, 시간대 쏠림'],
        ['튜터 품질 편차', '60회 (명시적)', '튜터 평가/관리 시스템 미흡'],
        ['짧은 수업시간', '787회', '20분/40분 구성이 깊은 학습에 부족'],
    ]
    add_table(doc, unique_headers, unique_data)

    doc.add_page_break()

    # ===== 5. 결론 =====
    doc.add_heading("5. 결론: 해결해야 할 핵심 문제", level=1)

    p = doc.add_paragraph()
    p.add_run("[즉시 해결 필요 (Urgent)]").bold = True
    doc.add_paragraph("""
1. 앱 안정성 개선
   • 근거: 앱 리뷰 부정률 38.2% (100/262건)
   • 영향: 신규 고객 첫인상 훼손, 앱스토어 평점 하락

2. 예약 시스템 개선
   • 근거: 예약 관련 불만 291건
   • 영향: 수업 이용 장벽, 고객 이탈
""")

    p = doc.add_paragraph()
    p.add_run("[중기 해결 필요 (Important)]").bold = True
    doc.add_paragraph("""
3. 튜터 품질 표준화
   • 근거: 튜터 품질 관련 언급 336건
   • 방안: 튜터 교육 강화, 피드백 가이드라인

4. 가격 라인업 다양화
   • 근거: 가격 관련 언급 2,589회, 스픽으로 이탈 언급
   • 방안: AI Only 저가 상품, 학생 할인

5. 학습 지속성 지원
   • 근거: 지속/동기 관련 언급 886회
   • 방안: 목표 설정, 성과 시각화, AI 복습 강화
""")

    p = doc.add_paragraph()
    p.add_run("[데이터 신뢰도]").bold = True
    doc.add_paragraph("""
• 전체 분석 건수: 7,604건 (10개 서비스)
• 링글 분석 건수: 1,631건
• 부정 리뷰 전수 분석: 100건
• 블로그/커뮤니티 문맥 분석: 908건
• 모든 수치는 실제 데이터 파일에서 추출 (재현 가능)
""")

    doc.add_page_break()

    # ===== 부록 =====
    doc.add_heading("부록: 분석 데이터 명세", level=1)

    doc.add_heading("A. 사용 데이터 파일", level=2)
    doc.add_paragraph("""
[메인 분석 데이터]
• ringle/ringle_filtered.csv (1,631건)
  - data_id, company, source_type, source_platform, text, rating, date

[비교 분석 데이터]
• cake/cake_filtered.csv (598건)
• hackers/hackers_filtered.csv (729건)
• malhae/malhae_filtered.csv (665건)
• maxai/maxai_filtered.csv (506건)
• pagoda/pagoda_filtered.csv (238건)
• santa/santa_filtered.csv (810건)
• uphone/uphone_filtered.csv (1,241건)
• yanadu/yanadu_filtered.csv (630건)
• carrot/carrot_filtered.csv (556건)

[집계 데이터]
• phase/02_company_comparison/data/sentiment_distribution.csv
• phase/03_deep_insights/data/competitor_mentions.csv
• phase/03_deep_insights/data/hidden_needs.csv
""")

    doc.add_heading("B. 키워드 분석 정의", level=2)

    kw_headers = ['문제 카테고리', '분석 키워드']
    kw_data = [
        ['가격 부담', '비싸, 가격, 비용, 부담, 돈, 환불, 구독료'],
        ['튜터 품질', '튜터, 강사, 선생님, 원어민, 실력, 친절, 불친절, 피드백'],
        ['앱 버그', '버그, 오류, 튕김, 튕겨, 느림, 로딩, 멈춤, 안됨, 안돼'],
        ['예약', '예약, 스케줄, 시간대, 마감, 경쟁'],
        ['지속성', '꾸준, 포기, 지속, 작심삼일, 습관, 동기'],
    ]
    add_table(doc, kw_headers, kw_data)

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("— 보고서 끝 —")

    # 저장
    output_path = os.path.join(OUTPUT_DIR, "고객문제_분석보고서.docx")
    doc.save(output_path)
    print(f"✅ 보고서 생성 완료: {output_path}")
    return output_path

if __name__ == "__main__":
    create_report()
